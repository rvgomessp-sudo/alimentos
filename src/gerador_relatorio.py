"""
Módulo Gerador de Relatório Forense em DOCX.

Gera documento Word formatado com análise forense completa,
incluindo gráficos matplotlib embutidos.
"""

import io
import json
from datetime import datetime
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MaxNLocator

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


RODAPE_TEXTO = "Holding Vazquez & Fonseca - Análise Técnica"

# Cores do tema
COR_PRIMARIA = RGBColor(0x1B, 0x3A, 0x5C)     # Azul escuro
COR_SECUNDARIA = RGBColor(0x2C, 0x5F, 0x8A)    # Azul médio
COR_DESTAQUE = RGBColor(0xC0, 0x39, 0x2B)       # Vermelho
COR_SUCESSO = RGBColor(0x27, 0xAE, 0x60)        # Verde
COR_TEXTO = RGBColor(0x2C, 0x3E, 0x50)          # Cinza escuro


def _estilizar_celula(celula, negrito=False, cor_fundo=None, alinhamento=None):
    """Aplica estilo a uma célula de tabela."""
    for paragrafo in celula.paragraphs:
        for run in paragrafo.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = COR_TEXTO
            if negrito:
                run.bold = True
        if alinhamento:
            paragrafo.alignment = alinhamento
    if cor_fundo:
        from docx.oxml.ns import qn
        shading = celula._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(
            qn("w:shd"),
            {qn("w:fill"): cor_fundo, qn("w:val"): "clear"},
        )
        shading.append(shading_elm)


def _grafico_para_bytes(fig):
    """Converte uma figura matplotlib em bytes PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


class GeradorRelatorioForense:
    """Gera relatório DOCX com análise forense processual completa."""

    def __init__(self, cruzamento, onerosidade, dados_processos=None):
        """
        Args:
            cruzamento: Resultado do CruzamentoForense.executar().
            onerosidade: Resultado da AnaliseOnerosidade.executar().
            dados_processos: Lista opcional de JSONs de análise individual.
        """
        self.cruzamento = cruzamento
        self.onerosidade = onerosidade
        self.dados_processos = dados_processos or []
        self.doc = Document()
        self._configurar_estilos()

    def _configurar_estilos(self):
        """Configura estilos base do documento."""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = COR_TEXTO

        for i in range(1, 5):
            nome = f"Heading {i}"
            if nome in self.doc.styles:
                h = self.doc.styles[nome]
                h.font.color.rgb = COR_PRIMARIA
                h.font.name = "Calibri"

        # Margens
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)

    def _adicionar_capa(self):
        """Gera a capa do relatório."""
        # Espaçamento superior
        for _ in range(6):
            self.doc.add_paragraph("")

        # Título principal
        titulo = self.doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titulo.add_run("RELATÓRIO DE ANÁLISE\nFORENSE PROCESSUAL")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = COR_PRIMARIA

        self.doc.add_paragraph("")

        # Linha separadora
        sep = self.doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep.add_run("━" * 40)
        run.font.color.rgb = COR_SECUNDARIA
        run.font.size = Pt(14)

        self.doc.add_paragraph("")

        # Processos analisados
        processos = self.cruzamento.get("processos_analisados", [])
        info = self.doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = info.add_run("Processos Analisados:\n")
        run.font.size = Pt(12)
        run.font.color.rgb = COR_SECUNDARIA

        run = info.add_run("1006347-84.2014.8.26.0020 (Execução de Alimentos)\n")
        run.font.size = Pt(11)
        run.font.color.rgb = COR_TEXTO

        run = info.add_run("0007068-67.2025.8.26.0020 (IDPJ)")
        run.font.size = Pt(11)
        run.font.color.rgb = COR_TEXTO

        self.doc.add_paragraph("")

        # Data
        data_p = self.doc.add_paragraph()
        data_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = data_p.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(12)
        run.font.color.rgb = COR_SECUNDARIA

        # Espaçamento inferior + rodapé
        for _ in range(4):
            self.doc.add_paragraph("")

        rodape = self.doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run(RODAPE_TEXTO)
        run.font.size = Pt(10)
        run.font.color.rgb = COR_SECUNDARIA
        run.italic = True

        self.doc.add_page_break()

    def _adicionar_identificacao(self):
        """Seção de identificação do caso."""
        self.doc.add_heading("IDENTIFICAÇÃO", level=1)

        # Tabela de identificação
        tabela = self.doc.add_table(rows=8, cols=2)
        tabela.style = "Light Grid Accent 1"
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER

        dados = [
            ("Processo Principal", "1006347-84.2014.8.26.0020"),
            ("Processo IDPJ", "0007068-67.2025.8.26.0020"),
            ("Vara", "5ª Vara de Família e Sucessões - N. Sra. do Ó - SP"),
            ("Exequente", "Nelson Wilson Fonseca Costa (CPF: 447.726.258-28)"),
            ("Executado", "Leandro Lopes da Costa (CPF: 259.171.178-02)"),
            ("Terceiro Interessado", "Jair Ribeiro da Costa (CPF: 614.786.528-53)"),
            ("Empresa (IDPJ)", "JRCIA Consultoria e Tecnologia Ltda (CNPJ: 01.950.077/0001-95)"),
            ("Período Analisado", self._periodo_analisado()),
        ]

        for i, (campo, valor) in enumerate(dados):
            tabela.rows[i].cells[0].text = campo
            tabela.rows[i].cells[1].text = valor
            _estilizar_celula(tabela.rows[i].cells[0], negrito=True, cor_fundo="D6E4F0")
            _estilizar_celula(tabela.rows[i].cells[1])

        self.doc.add_paragraph("")

    def _periodo_analisado(self):
        """Retorna o período analisado."""
        timeline = self.cruzamento.get("timeline_unificada", {})
        periodo = timeline.get("periodo", {})
        inicio = periodo.get("inicio", "N/D")
        fim = periodo.get("fim", "N/D")
        return f"{inicio} a {fim}"

    def _adicionar_resumo_executivo(self):
        """Seção 1: Resumo executivo com métricas principais."""
        self.doc.add_heading("1. RESUMO EXECUTIVO", level=1)

        # Score de correlação
        score = self.cruzamento.get("score_correlacao", {})
        score_val = score.get("score", 0)
        score_class = score.get("classificacao", "N/D")

        p = self.doc.add_paragraph()
        run = p.add_run(f"Score de Correlação entre Processos: {score_val}/100 ({score_class})")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COR_DESTAQUE if score_val >= 60 else COR_PRIMARIA

        self.doc.add_paragraph("")

        # Tabela de métricas
        tabela = self.doc.add_table(rows=1, cols=4)
        tabela.style = "Light Grid Accent 1"
        cabecalho = tabela.rows[0]
        for i, titulo in enumerate(["Métrica", "Exequente", "Executado", "Juízo"]):
            cabecalho.cells[i].text = titulo
            _estilizar_celula(cabecalho.cells[i], negrito=True, cor_fundo="1B3A5C")
            for run in cabecalho.cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        metricas = self.onerosidade.get("metricas_por_parte", {})
        linhas = [
            ("Total de Atos", "exequente", "executado", "juizo"),
            ("Petições", "exequente", "executado", "juizo"),
            ("Recursos", "exequente", "executado", "juizo"),
            ("Omissões >60 dias", "exequente", "executado", "juizo"),
        ]

        campos_map = {
            "Total de Atos": "total_atos",
            "Petições": "peticoes",
            "Recursos": "recursos",
            "Omissões >60 dias": "total_omissoes_60d",
        }

        for nome_metrica, *partes in linhas:
            row = tabela.add_row()
            row.cells[0].text = nome_metrica
            _estilizar_celula(row.cells[0], negrito=True, cor_fundo="EBF5FB")
            campo = campos_map[nome_metrica]
            for j, parte in enumerate(partes):
                val = metricas.get(parte, {}).get(campo, 0) or 0
                row.cells[j + 1].text = str(val)
                _estilizar_celula(row.cells[j + 1], alinhamento=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_paragraph("")

        # Índices
        protelacao = self.onerosidade.get("indice_protelacao", {})
        efetividade = self.onerosidade.get("indice_efetividade_exequente", {})

        p = self.doc.add_paragraph()
        run = p.add_run("Índice de Protelação do Executado: ")
        run.bold = True
        run = p.add_run(
            f"{protelacao.get('indice', 0):.2%} "
            f"({protelacao.get('classificacao', 'N/D')})"
        )
        run.font.color.rgb = COR_DESTAQUE

        p = self.doc.add_paragraph()
        run = p.add_run("Índice de Efetividade do Exequente: ")
        run.bold = True
        run = p.add_run(
            f"{efetividade.get('indice', 0):.2%} "
            f"({efetividade.get('classificacao', 'N/D')})"
        )

        # Entidades comuns
        entidades = self.cruzamento.get("entidades_comuns", {})
        cpfs_comuns = entidades.get("cpf", {}).get("total_comuns", 0)
        cnpjs_comuns = entidades.get("cnpj", {}).get("total_comuns", 0)

        self.doc.add_paragraph("")
        p = self.doc.add_paragraph()
        run = p.add_run("Entidades comuns entre os processos: ")
        run.bold = True
        run = p.add_run(f"{cpfs_comuns} CPFs, {cnpjs_comuns} CNPJs")

        self.doc.add_paragraph("")

    def _adicionar_timeline(self):
        """Seção 2: Timeline unificada dos processos."""
        self.doc.add_heading("2. TIMELINE UNIFICADA", level=1)

        timeline = self.cruzamento.get("timeline_unificada", {})
        eventos = timeline.get("eventos", [])

        p = self.doc.add_paragraph()
        run = p.add_run(f"Total de eventos na timeline: {timeline.get('total_eventos', 0)}")
        run.italic = True

        # Gráfico de timeline
        fig = self._gerar_grafico_timeline(eventos)
        if fig:
            buf = _grafico_para_bytes(fig)
            self.doc.add_picture(buf, width=Inches(6.2))
            ultimo_paragrafo = self.doc.paragraphs[-1]
            ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("")

        # Tabela com eventos principais (limitar a 60)
        eventos_relevantes = [
            e for e in eventos
            if e["tipo_evento"] in ("decisao", "sentenca", "constricao", "recurso", "movimentacao_societaria")
        ][:60]

        if eventos_relevantes:
            self.doc.add_heading("Eventos Relevantes", level=2)
            tabela = self.doc.add_table(rows=1, cols=4)
            tabela.style = "Light Grid Accent 1"

            for i, titulo in enumerate(["Data", "Processo", "Tipo", "Contexto"]):
                tabela.rows[0].cells[i].text = titulo
                _estilizar_celula(tabela.rows[0].cells[i], negrito=True, cor_fundo="1B3A5C")
                for run in tabela.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for ev in eventos_relevantes:
                row = tabela.add_row()
                row.cells[0].text = ev.get("data_formatada", "")
                row.cells[1].text = ev.get("processo", "")[:20]
                row.cells[2].text = ev.get("tipo_evento", "")
                row.cells[3].text = ev.get("contexto", "")[:80]
                for cell in row.cells:
                    _estilizar_celula(cell)

        self.doc.add_page_break()

    def _gerar_grafico_timeline(self, eventos):
        """Gera gráfico de dispersão temporal dos eventos por processo."""
        if not eventos:
            return None

        fig, ax = plt.subplots(figsize=(12, 4))

        cores_tipo = {
            "decisao": "#E74C3C",
            "sentenca": "#8E44AD",
            "constricao": "#E67E22",
            "recurso": "#3498DB",
            "peticao": "#2ECC71",
            "certidao": "#95A5A6",
            "intimacao": "#F39C12",
            "movimentacao_societaria": "#1ABC9C",
            "despacho": "#BDC3C7",
            "outros": "#D5D8DC",
        }

        processos_unicos = sorted(set(e["processo"] for e in eventos))
        y_map = {p: i for i, p in enumerate(processos_unicos)}

        for tipo, cor in cores_tipo.items():
            evs_tipo = [e for e in eventos if e["tipo_evento"] == tipo]
            if not evs_tipo:
                continue
            datas = []
            ys = []
            for e in evs_tipo:
                try:
                    dt = datetime.fromisoformat(e["data"])
                    datas.append(dt)
                    ys.append(y_map[e["processo"]])
                except (ValueError, TypeError):
                    continue
            if datas:
                ax.scatter(datas, ys, c=cor, label=tipo, alpha=0.6, s=15, edgecolors="none")

        ax.set_yticks(range(len(processos_unicos)))
        ax.set_yticklabels([p[:25] for p in processos_unicos], fontsize=8)
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%m/%Y"))
        ax.xaxis.set_major_locator(mdates.YearLocator())
        ax.set_xlabel("Data", fontsize=9)
        ax.set_title("Timeline Unificada dos Processos", fontsize=11, fontweight="bold")
        ax.legend(loc="upper left", fontsize=7, ncol=5, framealpha=0.8)
        ax.grid(axis="x", alpha=0.3)
        fig.tight_layout()

        return fig

    def _adicionar_onerosidade(self):
        """Seção 3: Análise de onerosidade com tabelas comparativas."""
        self.doc.add_heading("3. ANÁLISE DE ONEROSIDADE", level=1)

        # 3.1 Métricas comparativas
        self.doc.add_heading("3.1 Métricas por Parte", level=2)
        metricas = self.onerosidade.get("metricas_por_parte", {})

        tabela = self.doc.add_table(rows=1, cols=6)
        tabela.style = "Light Grid Accent 1"
        cabs = ["Parte", "Total Atos", "Petições", "Recursos", "Omissões >60d", "Tempo Médio (dias)"]
        for i, cab in enumerate(cabs):
            tabela.rows[0].cells[i].text = cab
            _estilizar_celula(tabela.rows[0].cells[i], negrito=True, cor_fundo="1B3A5C")
            for run in tabela.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for parte in ["exequente", "executado", "terceiro_jair", "juizo"]:
            m = metricas.get(parte, {})
            if not m or m.get("total_atos", 0) == 0:
                continue
            row = tabela.add_row()
            row.cells[0].text = m.get("nome", parte)[:35]
            row.cells[1].text = str(m.get("total_atos", 0))
            row.cells[2].text = str(m.get("peticoes", 0))
            row.cells[3].text = str(m.get("recursos", 0))
            row.cells[4].text = str(m.get("total_omissoes_60d", 0))
            tempo = m.get("tempo_medio_entre_atos_dias")
            row.cells[5].text = f"{tempo:.0f}" if tempo else "N/D"
            for cell in row.cells:
                _estilizar_celula(cell, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
            _estilizar_celula(row.cells[0], negrito=True, alinhamento=WD_ALIGN_PARAGRAPH.LEFT)

        self.doc.add_paragraph("")

        # Gráfico de barras comparativo
        fig = self._gerar_grafico_barras_comparativo(metricas)
        if fig:
            buf = _grafico_para_bytes(fig)
            self.doc.add_picture(buf, width=Inches(5.5))
            ultimo_paragrafo = self.doc.paragraphs[-1]
            ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph("")

        # 3.2 Índice de protelação
        self.doc.add_heading("3.2 Índice de Protelação do Executado", level=2)
        protelacao = self.onerosidade.get("indice_protelacao", {})

        p = self.doc.add_paragraph()
        run = p.add_run(f"Índice: {protelacao.get('indice', 0):.2%}")
        run.bold = True
        run.font.size = Pt(13)
        idx_val = protelacao.get("indice", 0)
        run.font.color.rgb = COR_DESTAQUE if idx_val >= 0.15 else COR_SUCESSO

        p = self.doc.add_paragraph()
        run = p.add_run(f"Classificação: {protelacao.get('classificacao', 'N/D')}")

        detalhes = protelacao.get("detalhes", {})
        p = self.doc.add_paragraph()
        p.add_run(f"Total de atos do executado: {detalhes.get('total_atos_executado', 0)}\n")
        p.add_run(f"Atos protelatórios identificados: {detalhes.get('atos_protelatarios', 0)}\n")
        p.add_run(f"Certidões negativas: {detalhes.get('certidoes_negativas', 0)}")

        self.doc.add_paragraph("")

        # 3.3 Efetividade do exequente
        self.doc.add_heading("3.3 Efetividade do Exequente", level=2)
        efet = self.onerosidade.get("indice_efetividade_exequente", {})

        p = self.doc.add_paragraph()
        run = p.add_run(f"Índice: {efet.get('indice', 0):.2%}")
        run.bold = True
        run.font.size = Pt(13)

        p = self.doc.add_paragraph()
        run = p.add_run(f"Classificação: {efet.get('classificacao', 'N/D')}")

        det = efet.get("detalhes", {})
        p = self.doc.add_paragraph()
        p.add_run(f"Tentativas de constrição: {det.get('total_tentativas_constricao', 0)}\n")
        p.add_run(f"Constrições exitosas: {det.get('constricoes_exitosas', 0)}\n")
        p.add_run(f"Constrições frustradas: {det.get('constricoes_frustradas', 0)}")

        self.doc.add_paragraph("")

        # 3.4 Heatmap de atividade mensal
        self.doc.add_heading("3.4 Atividade Mensal", level=2)

        dados_graf = self.onerosidade.get("dados_graficos", {})
        fig = self._gerar_heatmap_atividade(dados_graf.get("atividade_mensal", {}))
        if fig:
            buf = _grafico_para_bytes(fig)
            self.doc.add_picture(buf, width=Inches(6.2))
            ultimo_paragrafo = self.doc.paragraphs[-1]
            ultimo_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_page_break()

    def _gerar_grafico_barras_comparativo(self, metricas):
        """Gera gráfico de barras comparativo entre partes."""
        partes_labels = []
        total_atos = []
        peticoes = []
        recursos = []

        for parte in ["exequente", "executado", "juizo"]:
            m = metricas.get(parte, {})
            if m.get("total_atos", 0) > 0:
                partes_labels.append(m.get("nome", parte)[:20])
                total_atos.append(m.get("total_atos", 0))
                peticoes.append(m.get("peticoes", 0))
                recursos.append(m.get("recursos", 0))

        if not partes_labels:
            return None

        fig, ax = plt.subplots(figsize=(8, 4))
        x = range(len(partes_labels))
        largura = 0.25

        ax.bar([i - largura for i in x], total_atos, largura, label="Total Atos", color="#2C5F8A")
        ax.bar(x, peticoes, largura, label="Petições", color="#27AE60")
        ax.bar([i + largura for i in x], recursos, largura, label="Recursos", color="#E74C3C")

        ax.set_xticks(x)
        ax.set_xticklabels(partes_labels, fontsize=8)
        ax.set_ylabel("Quantidade", fontsize=9)
        ax.set_title("Comparativo de Atos por Parte", fontsize=11, fontweight="bold")
        ax.legend(fontsize=8)
        ax.yaxis.set_major_locator(MaxNLocator(integer=True))
        ax.grid(axis="y", alpha=0.3)
        fig.tight_layout()

        return fig

    def _gerar_heatmap_atividade(self, atividade_mensal):
        """Gera heatmap de atividade mensal por parte."""
        if not atividade_mensal:
            return None

        meses = sorted(atividade_mensal.keys())
        partes = ["executado", "exequente", "juizo"]
        partes_labels = ["Executado", "Exequente", "Juízo"]

        # Montar matriz
        matriz = []
        for parte in partes:
            linha = [atividade_mensal.get(mes, {}).get(parte, 0) for mes in meses]
            matriz.append(linha)

        if not any(any(v > 0 for v in linha) for linha in matriz):
            return None

        fig, ax = plt.subplots(figsize=(max(12, len(meses) * 0.4), 3))

        im = ax.imshow(matriz, aspect="auto", cmap="YlOrRd", interpolation="nearest")

        ax.set_yticks(range(len(partes_labels)))
        ax.set_yticklabels(partes_labels, fontsize=8)

        # Mostrar apenas alguns meses para legibilidade
        step = max(1, len(meses) // 20)
        tick_positions = list(range(0, len(meses), step))
        ax.set_xticks(tick_positions)
        ax.set_xticklabels([meses[i] for i in tick_positions], rotation=45, ha="right", fontsize=7)

        ax.set_title("Heatmap de Atividade Mensal", fontsize=11, fontweight="bold")
        fig.colorbar(im, ax=ax, shrink=0.8, label="Eventos")
        fig.tight_layout()

        return fig

    def _adicionar_correlacoes(self):
        """Seção 4: Correlações detectadas entre processos."""
        self.doc.add_heading("4. CORRELAÇÕES ENTRE PROCESSOS", level=1)

        # Score
        score = self.cruzamento.get("score_correlacao", {})
        componentes = score.get("componentes", {})

        self.doc.add_heading("4.1 Score de Correlação", level=2)

        tabela = self.doc.add_table(rows=1, cols=3)
        tabela.style = "Light Grid Accent 1"
        for i, cab in enumerate(["Componente", "Score", "Peso"]):
            tabela.rows[0].cells[i].text = cab
            _estilizar_celula(tabela.rows[0].cells[i], negrito=True, cor_fundo="1B3A5C")
            for run in tabela.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        nomes_componentes = {
            "entidades_comuns": "Entidades Comuns",
            "proximidade_temporal": "Proximidade Temporal",
            "padroes_suspeitos": "Padrões Suspeitos",
        }

        for chave, nome in nomes_componentes.items():
            comp = componentes.get(chave, {})
            row = tabela.add_row()
            row.cells[0].text = nome
            row.cells[1].text = f"{comp.get('score', 0):.1f}/100"
            row.cells[2].text = comp.get("peso", "")
            for cell in row.cells:
                _estilizar_celula(cell, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)

        row = tabela.add_row()
        row.cells[0].text = "SCORE FINAL"
        row.cells[1].text = f"{score.get('score', 0):.1f}/100 ({score.get('classificacao', '')})"
        row.cells[2].text = "100%"
        for cell in row.cells:
            _estilizar_celula(cell, negrito=True, cor_fundo="D6E4F0", alinhamento=WD_ALIGN_PARAGRAPH.CENTER)

        self.doc.add_paragraph("")

        # Entidades comuns
        self.doc.add_heading("4.2 Entidades Comuns", level=2)
        entidades = self.cruzamento.get("entidades_comuns", {})

        for tipo in ["cpf", "cnpj"]:
            dados_tipo = entidades.get(tipo, {})
            identificados = dados_tipo.get("identificados", [])
            if identificados:
                p = self.doc.add_paragraph()
                run = p.add_run(f"{tipo.upper()}s comuns ({dados_tipo.get('total_comuns', 0)}):")
                run.bold = True

                for item in identificados:
                    p = self.doc.add_paragraph(style="List Bullet")
                    valor = item.get(tipo, "")
                    nome = item.get("nome", "")
                    papel = item.get("papel", "")
                    p.add_run(f"{valor} - {nome}")
                    if papel:
                        run = p.add_run(f" ({papel})")
                        run.italic = True

        self.doc.add_paragraph("")

        # Padrões suspeitos
        padroes = self.cruzamento.get("padroes_suspeitos", {})

        self.doc.add_heading("4.3 Reações Pós-Decisão Desfavorável", level=2)
        reacoes = padroes.get("reacoes_pos_decisao", {})
        p = self.doc.add_paragraph()
        p.add_run(
            f"Decisões desfavoráveis ao executado: {reacoes.get('total_decisoes_desfavoraveis', 0)}\n"
        )
        p.add_run(
            f"Reações do executado em até {reacoes.get('janela_dias', 30)} dias: "
            f"{reacoes.get('correlacoes_encontradas', 0)}"
        )

        detalhes = reacoes.get("detalhes", [])[:10]
        if detalhes:
            tabela = self.doc.add_table(rows=1, cols=4)
            tabela.style = "Light Grid Accent 1"
            for i, cab in enumerate(["Decisão (data)", "Reação (data)", "Dias", "Contexto Reação"]):
                tabela.rows[0].cells[i].text = cab
                _estilizar_celula(tabela.rows[0].cells[i], negrito=True, cor_fundo="C0392B")
                for run in tabela.rows[0].cells[i].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            for d in detalhes:
                row = tabela.add_row()
                row.cells[0].text = d.get("decisao", {}).get("data", "")
                row.cells[1].text = d.get("reacao", {}).get("data", "")
                row.cells[2].text = str(d.get("dias_entre", ""))
                row.cells[3].text = d.get("reacao", {}).get("contexto", "")[:60]
                for cell in row.cells:
                    _estilizar_celula(cell)

        self.doc.add_paragraph("")

        self.doc.add_heading("4.4 Movimentações Societárias Coincidentes", level=2)
        mov = padroes.get("movimentacoes_societarias", {})
        p = self.doc.add_paragraph()
        p.add_run(
            f"Movimentações societárias (JRCIA): {mov.get('total_movimentacoes_societarias', 0)}\n"
        )
        p.add_run(
            f"Coincidências com atos de execução (90 dias): {mov.get('coincidencias_90_dias', 0)}"
        )

        self.doc.add_paragraph("")

        self.doc.add_heading("4.5 Padrões de Atividade", level=2)
        atividade = padroes.get("padroes_atividade", {})
        exec_data = atividade.get("executado", {})

        p = self.doc.add_paragraph()
        p.add_run(f"Eventos do executado: {exec_data.get('total_eventos', 0)}\n")
        p.add_run(f"Períodos de omissão (>60 dias): {exec_data.get('periodos_omissao_60d', 0)}\n")
        p.add_run(f"Períodos de alta atividade: {exec_data.get('periodos_alta_atividade', 0)}")

        omissoes = exec_data.get("detalhes_omissao", [])[:5]
        if omissoes:
            p = self.doc.add_paragraph()
            run = p.add_run("Maiores períodos de omissão:")
            run.bold = True
            for om in omissoes:
                p = self.doc.add_paragraph(style="List Bullet")
                p.add_run(f"{om['de']} a {om['ate']} ({om['dias']} dias)")

        self.doc.add_page_break()

    def _adicionar_conclusoes(self):
        """Seção 5: Conclusões técnicas."""
        self.doc.add_heading("5. CONCLUSÕES TÉCNICAS", level=1)

        score = self.cruzamento.get("score_correlacao", {})
        protelacao = self.onerosidade.get("indice_protelacao", {})
        efetividade = self.onerosidade.get("indice_efetividade_exequente", {})
        padroes = self.cruzamento.get("padroes_suspeitos", {})
        entidades = self.cruzamento.get("entidades_comuns", {})

        conclusoes = []

        # 1. Correlação
        score_val = score.get("score", 0)
        if score_val >= 60:
            conclusoes.append(
                f"A análise quantitativa revela correlação {score.get('classificacao', '').lower()} "
                f"(score {score_val}/100) entre os processos de execução de alimentos e o IDPJ, "
                f"indicando interdependência significativa entre os feitos."
            )

        # 2. Entidades comuns
        cpfs_comuns = entidades.get("cpf", {}).get("total_comuns", 0)
        cnpjs_comuns = entidades.get("cnpj", {}).get("total_comuns", 0)
        if cpfs_comuns > 0 or cnpjs_comuns > 0:
            conclusoes.append(
                f"Foram identificados {cpfs_comuns} CPFs e {cnpjs_comuns} CNPJs comuns aos dois processos, "
                f"confirmando a vinculação subjetiva entre as partes e a empresa JRCIA Consultoria e "
                f"Tecnologia Ltda (CNPJ 01.950.077/0001-95)."
            )

        # 3. Protelação
        idx_prot = protelacao.get("indice", 0)
        if idx_prot > 0:
            conclusoes.append(
                f"O índice de protelação do executado é de {idx_prot:.2%} "
                f"({protelacao.get('classificacao', '')}), "
                f"com {protelacao.get('detalhes', {}).get('atos_protelatarios', 0)} atos "
                f"protelatórios identificados ao longo do processo."
            )

        # 4. Efetividade
        idx_efet = efetividade.get("indice", 0)
        conclusoes.append(
            f"A efetividade das tentativas de constrição patrimonial pelo exequente é "
            f"{efetividade.get('classificacao', '').lower()} ({idx_efet:.2%}), "
            f"com {efetividade.get('detalhes', {}).get('constricoes_frustradas', 0)} "
            f"tentativas frustradas de um total de "
            f"{efetividade.get('detalhes', {}).get('total_tentativas_constricao', 0)}."
        )

        # 5. Padrões
        reacoes = padroes.get("reacoes_pos_decisao", {})
        if reacoes.get("correlacoes_encontradas", 0) > 0:
            conclusoes.append(
                f"Foram detectadas {reacoes['correlacoes_encontradas']} correlações temporais "
                f"entre decisões desfavoráveis e reações do executado dentro de uma janela "
                f"de {reacoes.get('janela_dias', 30)} dias, sugerindo padrão reativo."
            )

        # 6. Omissões
        exec_ativ = padroes.get("padroes_atividade", {}).get("executado", {})
        omissoes = exec_ativ.get("periodos_omissao_60d", 0)
        if omissoes > 0:
            conclusoes.append(
                f"O executado apresentou {omissoes} períodos de omissão superiores a 60 dias, "
                f"intercalados com períodos de alta atividade processual, padrão compatível "
                f"com estratégia de procrastinação seletiva."
            )

        # Escrever conclusões
        for i, conclusao in enumerate(conclusoes, 1):
            p = self.doc.add_paragraph()
            run = p.add_run(f"{i}. ")
            run.bold = True
            p.add_run(conclusao)
            self.doc.add_paragraph("")

        # Rodapé final
        self.doc.add_paragraph("")
        sep = self.doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep.add_run("━" * 40)
        run.font.color.rgb = COR_SECUNDARIA

        rodape = self.doc.add_paragraph()
        rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = rodape.add_run(RODAPE_TEXTO)
        run.font.size = Pt(10)
        run.font.color.rgb = COR_SECUNDARIA
        run.italic = True

        data_p = self.doc.add_paragraph()
        data_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = data_p.add_run(f"Relatório gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        run.font.size = Pt(9)
        run.font.color.rgb = COR_SECUNDARIA

    def gerar(self, caminho_saida):
        """
        Gera o relatório DOCX completo.

        Args:
            caminho_saida: Caminho para salvar o arquivo .docx.
        """
        self._adicionar_capa()
        self._adicionar_identificacao()
        self._adicionar_resumo_executivo()
        self._adicionar_timeline()
        self._adicionar_onerosidade()
        self._adicionar_correlacoes()
        self._adicionar_conclusoes()

        caminho = Path(caminho_saida)
        caminho.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(caminho))

        return str(caminho)
